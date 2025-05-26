from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from .models import Deal, Buyer, Realtor, UserProfile
from .forms import DealForm, BuyerForm, RealtorForm, UserRegistrationForm, UserProfileForm
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from django.http import HttpResponse

# Deal views
def deal_list(request):
    deals = Deal.objects.all().order_by('-deal_date')
    return render(request, 'agency/deal_list.html', {'deals': deals})

@login_required
def deal_create(request):
    if request.method == 'POST':
        form = DealForm(request.POST)
        if form.is_valid():
            deal = form.save(commit=False)
            deal.created_by = request.user
            deal.save()
            return render(request, 'agency/success.html', {'message': 'Сделка успешно добавлена!'})
    else:
        form = DealForm()
    return render(request, 'agency/deal_form.html', {'form': form})

def deal_detail(request, pk):
    deal = get_object_or_404(Deal, pk=pk)
    return render(request, 'agency/deal_detail.html', {'deal': deal})

@login_required
def deal_edit(request, pk):
    deal = get_object_or_404(Deal, pk=pk)
    if request.method == "POST":
        form = DealForm(request.POST, instance=deal)
        if form.is_valid():
            form.save()
            return redirect('agency:deal_detail', pk=deal.pk)
    else:
        form = DealForm(instance=deal)
    return render(request, 'agency/deal_form.html', {'form': form, 'deal': deal})

@login_required
def deal_delete(request, pk):
    deal = get_object_or_404(Deal, pk=pk)
    if request.method == 'POST':
        deal.delete()
        return redirect('agency:deal_list')
    return render(request, 'agency/deal_confirm_delete.html', {'deal': deal})

# Buyer views
def buyer_list(request):
    buyers = Buyer.objects.all().order_by('last_name')
    return render(request, 'agency/buyer_list.html', {'buyers': buyers})

@login_required
def buyer_create(request):
    if request.method == 'POST':
        form = BuyerForm(request.POST)
        if form.is_valid():
            buyer = form.save(commit=False)
            buyer.created_by = request.user
            buyer.save()
            return render(request, 'agency/success.html', {'message': 'Покупатель успешно добавлен!'})
    else:
        form = BuyerForm()
    return render(request, 'agency/buyer_form.html', {'form': form})

def buyer_detail(request, pk):
    buyer = get_object_or_404(Buyer, pk=pk)
    return render(request, 'agency/buyer_detail.html', {'buyer': buyer})

@login_required
def buyer_edit(request, pk):
    buyer = get_object_or_404(Buyer, pk=pk)
    if request.method == "POST":
        form = BuyerForm(request.POST, instance=buyer)
        if form.is_valid():
            form.save()
            return redirect('agency:buyer_detail', pk=buyer.pk)
    else:
        form = BuyerForm(instance=buyer)
    return render(request, 'agency/buyer_form.html', {'form': form, 'buyer': buyer})

@login_required
def buyer_delete(request, pk):
    buyer = get_object_or_404(Buyer, pk=pk)
    if request.method == 'POST':
        buyer.delete()
        return redirect('agency:buyer_list')
    return render(request, 'agency/buyer_confirm_delete.html', {'buyer': buyer})

# Realtor views
def realtor_list(request):
    realtors = Realtor.objects.all().order_by('last_name')
    return render(request, 'agency/realtor_list.html', {'realtors': realtors})

@login_required
def realtor_create(request):
    if request.method == 'POST':
        form = RealtorForm(request.POST)
        if form.is_valid():
            realtor = form.save(commit=False)
            realtor.created_by = request.user
            realtor.save()
            return render(request, 'agency/success.html', {'message': 'Риелтор успешно добавлен!'})
    else:
        form = RealtorForm()
    return render(request, 'agency/realtor_form.html', {'form': form})

def realtor_detail(request, pk):
    realtor = get_object_or_404(Realtor, pk=pk)
    return render(request, 'agency/realtor_detail.html', {'realtor': realtor})

@login_required
def realtor_edit(request, pk):
    realtor = get_object_or_404(Realtor, pk=pk)
    if request.method == "POST":
        form = RealtorForm(request.POST, instance=realtor)
        if form.is_valid():
            form.save()
            return redirect('agency:realtor_detail', pk=realtor.pk)
    else:
        form = RealtorForm(instance=realtor)
    return render(request, 'agency/realtor_form.html', {'form': form, 'realtor': realtor})

@login_required
def realtor_delete(request, pk):
    realtor = get_object_or_404(Realtor, pk=pk)
    if request.method == 'POST':
        realtor.delete()
        return redirect('agency:realtor_list')
    return render(request, 'agency/realtor_confirm_delete.html', {'realtor': realtor})

# Authentication views
def user_login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('agency:deal_list')
        else:
            return render(request, 'agency/login.html', {'error': 'Неверное имя пользователя или пароль.'})
    return render(request, 'agency/login.html')

def user_logout(request):
    logout(request)
    return redirect('agency:login')

def user_register(request):
    print("Entering user_register")  # Для отладки
    if request.method == 'POST':
        print("POST request received")
        user_form = UserRegistrationForm(request.POST)
        profile_form = UserProfileForm(request.POST)
        if user_form.is_valid() and profile_form.is_valid():
            user = user_form.save(commit=False)
            user.set_password(user_form.cleaned_data['password'])
            user.save()
            profile = profile_form.save(commit=False)
            profile.user = user
            profile.save()
            login(request, user)
            return render(request, 'agency/success.html', {'message': 'Регистрация успешна!'})
    else:
        print("GET request received")
        user_form = UserRegistrationForm()
        profile_form = UserProfileForm()
    return render(request, 'agency/register.html', {'user_form': user_form, 'profile_form': profile_form})

# Excel экспорт
def export_deals_excel(request):
    deals = Deal.objects.all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Сделки"

    # Заголовки
    ws.append(['ID', 'Покупатель', 'Риелтор', 'Адрес', 'Цена', 'Дата сделки'])

    for deal in deals:
        ws.append([
            deal.id,
            str(deal.buyer),
            str(deal.realtor),
            deal.property_address,
            deal.price,
            deal.deal_date.strftime("%Y-%m-%d") if deal.deal_date else ""
        ])

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = 'attachment; filename=deals.xlsx'
    wb.save(response)
    return response

# Word экспорт
def export_deals_word(request):
    deals = Deal.objects.all()
    doc = Document()
    doc.add_heading('Список сделок', 0)

    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Покупатель'
    hdr_cells[2].text = 'Риелтор'
    hdr_cells[3].text = 'Адрес'
    hdr_cells[4].text = 'Цена'
    hdr_cells[5].text = 'Дата сделки'

    for deal in deals:
        row_cells = table.add_row().cells
        row_cells[0].text = str(deal.id)
        row_cells[1].text = str(deal.buyer)
        row_cells[2].text = str(deal.realtor)
        row_cells[3].text = deal.property_address
        row_cells[4].text = str(deal.price)
        row_cells[5].text = deal.deal_date.strftime("%Y-%m-%d") if deal.deal_date else ""

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=deals.docx'
    return response

# PDF экспорт
def export_deals_pdf(request):
    deals = Deal.objects.all()
    buffer = BytesIO()
    p = canvas.Canvas(buffer)

    y = 800
    p.setFont("Helvetica-Bold", 14)
    p.drawString(100, y, "Список сделок")
    y -= 30

    p.setFont("Helvetica-Bold", 10)
    headers = ['ID', 'Покупатель', 'Риелтор', 'Адрес', 'Цена', 'Дата сделки']
    x_positions = [30, 60, 160, 260, 460, 510]
    for i, header in enumerate(headers):
        p.drawString(x_positions[i], y, header)

    y -= 20
    p.setFont("Helvetica", 9)
    for deal in deals:
        if y < 40:
            p.showPage()
            y = 800
        p.drawString(x_positions[0], y, str(deal.id))
        p.drawString(x_positions[1], y, str(deal.buyer))
        p.drawString(x_positions[2], y, str(deal.realtor))
        p.drawString(x_positions[3], y, deal.property_address)
        p.drawString(x_positions[4], y, str(deal.price))
        p.drawString(x_positions[5], y, deal.deal_date.strftime("%Y-%m-%d") if deal.deal_date else "")
        y -= 20

    p.save()
    pdf = buffer.getvalue()
    buffer.close()

    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=deals.pdf'
    return response